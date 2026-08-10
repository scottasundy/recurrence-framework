#!/usr/bin/env python3
"""Build the the Cosmological Recurrence Study manuscript Markdown and DOCX from reproduced outputs."""
from __future__ import annotations

from pathlib import Path
import math
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
FIG = ROOT / "figures"
PAPER = ROOT / "paper"
PAPER.mkdir(exist_ok=True)

baseline = pd.read_csv(OUT / "de-sitter-baseline.csv").iloc[0]
post = pd.read_csv(OUT / "de-sitter-posterior.csv").set_index("quantity")
gauss = pd.read_csv(OUT / "gaussian-check.csv").set_index("quantity")
prior = pd.read_csv(OUT / "prior-sensitivity.csv")
bigrip = pd.read_csv(OUT / "big-rip.csv")
model_weights = pd.read_csv(OUT / "model-weights.csv")
fate = pd.read_csv(OUT / "fate-sensitivity.csv")


def sci(x, digits=4):
    if x == 0:
        return "0"
    e = int(math.floor(math.log10(abs(float(x)))))
    m = float(x) / 10**e
    return f"{m:.{digits}f} x 10^{e}"


def set_cell_shading(cell, fill="EAECEF"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


md = []
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(24)
styles["Heading 1"].font.name = "Aptos Display"
styles["Heading 1"].font.size = Pt(16)
styles["Heading 2"].font.name = "Aptos Display"
styles["Heading 2"].font.size = Pt(12.5)


def add_heading(text, level=1):
    doc.add_heading(text, level=level)
    md.append("#" * level + " " + text + "\n")


def add_para(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    md.append(text + "\n")


def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")
    md.append(f"- {text}\n")


def add_equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    md.append(f"$$ {text} $$\n")


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = str(h)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c)
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    md.append("| " + " | ".join(map(str, headers)) + " |\n")
    md.append("|" + "|".join(["---"] * len(headers)) + "|\n")
    for row in rows:
        md.append("| " + " | ".join(str(x).replace("|", "\\|") for x in row) + " |\n")
    md.append("\n")
    return t


def add_figure(filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG / filename), width=Inches(6.4))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
    md.append(f"![{caption}](../figures/{filename})\n\n")


# Cover
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Cosmological Recurrence Probability Study")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Fate-Conditioned Partial-Identification Framework for Cosmological Recurrence")
r.bold = True
r.font.size = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("The study - Publication Release")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Scott A. Sundy\n9 August 2026")
md.extend([
    "# Cosmological Recurrence Probability Study\n",
    "**A Fate-Conditioned Partial-Identification Framework for Cosmological Recurrence**\n",
    "The study - Publication Release  \nScott A. Sundy  \n9 August 2026\n",
])
doc.add_page_break()

add_heading("Research question", 1)
add_para("Given the universe we observe today, which physically viable futures permit recurrence, what timescale follows under each declared recurrence model, and what probability can be assigned after accounting for uncertainty in the universe's ultimate fate?")

add_heading("Abstract", 1)
add_para(
    "Cosmological recurrence is often reduced to a single spectacular timescale or an unsupported probability. This study instead treats it as a structured inference problem. It separates observational parameter uncertainty inside a specified cosmological model from uncertainty about the far-future cosmological branch and from additional quantum assumptions required for recurrence. It also distinguishes exact microscopic recurrence, near recurrence, causal-patch recurrence, macrostate recurrence, and observational recurrence."
)
add_para(
    f"For an eternal, stable, spatially flat positive-Lambda future, the Cosmic Coordinate reference point gives an asymptotic de Sitter horizon entropy S_dS/k_B = {sci(baseline['S_dS_over_kB'], 5)}. The conventional finite-patch thermodynamic estimate t_thermo ~ H_Lambda^-1 exp(S_dS/k_B) therefore gives log10(t_thermo/yr) = {sci(baseline['log10_t_thermo_rec_years'], 5)}. This quantity is a conditional entropy-exponential thermodynamic recurrence scale. It is not a computed epsilon-recurrence time for the observed universe."
)
add_para(
    f"The primary observational propagation uses the released DESI DR2+CMB flat-LambdaCDM Cobaya posterior rather than a Gaussian surrogate. The weighted posterior median is log10(t_thermo/yr) = {sci(post.loc['log10_t_thermo_rec_years','median'], 5)}, with a 95% interval [{sci(post.loc['log10_t_thermo_rec_years','q2.5'], 5)}, {sci(post.loc['log10_t_thermo_rec_years','q97.5'], 5)}]. This observational spread is negligible relative to uncertainty in the universe's ultimate fate and in the microscopic assumptions behind recurrence."
)
add_para(
    f"The study retains a data-informed model-family layer using the Ong, Yallup & Handley Bayesian reanalysis. The DESI DR2 BAO + Planck CMB baseline is ln B_dynamic/Lambda = -0.57 +/- 0.26; the corrected DES-Dovekie comparison is updated to the current reported value -0.01 +/- 0.27. With equal prior odds restricted to LambdaCDM and w0-wa CDM, the baseline corresponds to posterior model weights of {model_weights.iloc[0].p_LambdaCDM*100:.1f}% for LambdaCDM and {model_weights.iloc[0].p_dynamic_DE*100:.1f}% for dynamical dark energy. These are probabilities over a declared two-model set, not probabilities of specific ultimate fates."
)
add_para(
    "The principal probability result is partial identification. Present observations do not determine the probability that dark energy remains a positive cosmological constant forever, the lifetime of our vacuum relative to recurrence, whether a de Sitter causal patch is described by a finite effectively closed Hilbert space, or whether any future bounce has recurrent dynamics. Without theoretical priors on those unknowns, the marginalized recurrence probability remains [0,1]. Because this is the full logical probability range, it is a non-identifiability result rather than an informative probability estimate. Narrower intervals are therefore reported only as prior-conditioned theoretical sensitivity analyses, never as measured cosmological probabilities."
)

doc.add_page_break()
add_heading("Key results", 2)
add_table(
    ["Quantity", "Result", "Interpretation"],
    [
        ["Cosmic Coordinate de Sitter horizon time", f"{baseline['horizon_time_Gyr']:.4f} Gyr", "Model-conditional reference"],
        ["Cosmic Coordinate de Sitter entropy", sci(baseline['S_dS_over_kB'], 5) + " k_B", "Model-conditional reference"],
        ["CPS thermodynamic recurrence exponent", sci(baseline['log10_t_thermo_rec_years'], 5), "Conditional thermodynamic scale; not epsilon recurrence"],
        ["DESI DR2+CMB chain median exponent", sci(post.loc['log10_t_thermo_rec_years','median'], 5), "Weighted official posterior propagation under flat LambdaCDM"],
        ["Equal-prior two-model weight: LambdaCDM", f"{model_weights.iloc[0].p_LambdaCDM*100:.1f}%", "Published Bayesian evidence; only LambdaCDM vs w0-wa CDM"],
        ["Equal-prior two-model weight: dynamical DE", f"{model_weights.iloc[0].p_dynamic_DE*100:.1f}%", "Published Bayesian evidence; not an ultimate-fate probability"],
        ["Data-only marginalized recurrence probability", "[0,1]", "Ultimate-fate/theory probabilities are not identified"],
        ["Quantum epsilon recurrence", "Not numerically identified for cosmology", "Requires a finite discrete spectrum, energy span, trace distance, and epsilon"],
    ],
)

add_heading("1. Scope and scientific status", 1)
add_para("CRPS is a standalone cosmological study. Cosmic Coordinate is used only as a declared reference point for one flat matter+Lambda calculation. Recurrence Dynamics is not used to prove CRPS and is not incorporated into the cosmological inference. Its earlier exact-versus-near terminology is retained only where it improves clarity.")
add_para("The principal observational input is the official DESI DR2 Results II flat-LambdaCDM posterior for all DESI DR2 BAO measurements combined with the default CMB likelihood set. The archive includes a compact projection of the four released Cobaya chains to the only fields required here: posterior weight, H0, and Omega_m. The source files are independently hash-verified against DESI's published checksum catalog.")
add_para("DESI DR2 Results IV, released in July 2026, adds a Lyman-alpha Alcock-Paczynski/full-shape measurement and reports updated extended-model constraints. It strengthens the reason not to treat flat LambdaCDM as established far-future physics. CRPS therefore uses flat LambdaCDM only as a conditional branch for calculating the de Sitter horizon scale; it does not claim that current data establish an eternal cosmological constant.")

doc.add_page_break()
add_heading("2. Recurrence targets must be declared", 1)
add_para("The statement 'the universe recurs' is incomplete until the object being compared, the distance measure, and the tolerance are specified. CRPS separates the following targets.")
add_table(
    ["Target", "Definition", "Status in CRPS"],
    [
        ["Exact microscopic recurrence", "Complete microscopic/quantum state returns exactly", "Not inferred from cosmological data"],
        ["Quantum near recurrence", "Finite-system states return within trace distance epsilon", "Theorem available; cosmological spectrum unknown"],
        ["Causal-patch thermodynamic recurrence", "Entropy-exponential fluctuation/return scale of an effectively finite de Sitter patch", "Conditional calculable scale"],
        ["Cosmological macrostate recurrence", "Declared macroscopic variables return to a target region", "Requires an explicit macrostate metric"],
        ["Observational recurrence", "Selected observables become indistinguishable within stated resolution", "Requires a declared observable set and tolerances"],
    ],
)
add_equation("R_epsilon(t): d[x(t), x(0)] <= epsilon")
add_para("The study preserves the separation between the de Sitter thermodynamic clock and quantum epsilon recurrence, and adds a hierarchical fate-probability layer: observed data constrain top-level model-family weights, while unresolved within-family fate splits are exposed as explicit theory sensitivity assumptions rather than hidden inside a single percentage.")

add_heading("3. Fate-conditioned probability architecture", 1)
add_para("Let D denote present observations, M a far-future cosmological model, theta measured cosmological parameters within M, and A additional theoretical assumptions required for a recurrence theorem or mechanism. The target probability has the schematic form:")
add_equation("P(R_epsilon < T | D) = sum_{M,A} integral P(R_epsilon < T | M, theta, A) p(theta, M, A | D) dtheta")
add_para("Only part of this expression is presently data-driven. The posterior p(theta | M,D) can be propagated from released cosmological chains. The probabilities of the far-future models and quantum assumptions are not presently identified. CRPS therefore refuses to collapse them into a single measured percentage.")
add_bullet("Observational uncertainty: H0, Omega_m, and other parameters inside a declared cosmological model.")
add_bullet("Future-model uncertainty: eternal Lambda, metastable vacuum, fading dark energy, finite-time endpoint, recollapse, or cyclic/bounce behavior.")
add_bullet("Recurrence-theory uncertainty: finite effective state space, unitarity/closure, spectral structure, vacuum survival, and recurrent properties of any cycle map.")

add_heading("4. Physically distinct future branches", 1)
add_table(
    ["Future branch", "Recurrence status", "What must additionally be true"],
    [
        ["Eternal positive Lambda / de Sitter", "Permits a conditional thermodynamic recurrence argument", "Finite/effectively closed recurrent patch interpretation"],
        ["Metastable de Sitter", "Possible only before vacuum decay", "Vacuum survival must compete successfully with recurrence"],
        ["Phantom finite-lifetime endpoint", "No asymptotic de Sitter recurrence regime", "Future really remains sufficiently phantom"],
        ["Dark energy fades toward zero", "No standard finite-patch guarantee", "Asymptotic state-space/horizon structure must be specified"],
        ["Recollapse without bounce", "Contraction alone does not imply recurrence", "A return map is required"],
        ["Cyclic/bouncing future", "Potentially recurrent", "Cycle map itself must return or approach prior states"],
        ["Spatial duplicates / eternal inflation", "Excluded from target", "A duplicate elsewhere is not temporal recurrence of this patch"],
    ],
)

add_heading("5. Eternal de Sitter thermodynamic branch", 1)
add_heading("5.1 Horizon scale and entropy", 2)
add_equation("H_Lambda = H0 sqrt(Omega_Lambda)")
add_equation("r_dS = c / H_Lambda")
add_equation("S_dS/k_B = pi r_dS^2 / l_P^2")
add_para(
    f"At the Cosmic Coordinate reference point H0 = 68.11 km s^-1 Mpc^-1 and Omega_Lambda = 0.6958, CRPS obtains H_Lambda^-1 = {baseline['horizon_time_Gyr']:.4f} Gyr, r_dS = {sci(baseline['horizon_radius_m'],4)} m, and S_dS/k_B = {sci(baseline['S_dS_over_kB'],5)}."
)
add_heading("5.2 Conditional entropy-exponential recurrence scale", 2)
add_equation("t_thermo ~ H_Lambda^-1 exp(S_dS/k_B)")
add_para(
    f"The scale is too large to construct directly. In logarithmic form, log10(t_thermo/yr) = {sci(baseline['log10_t_thermo_rec_years'],5)} and log10(log10(t_thermo/yr)) = {baseline['log10log10_t_thermo_rec_years']:.5f}. The latter is only a compact plotting coordinate."
)
add_para("This is a conventional thermodynamic/Poincare estimate under a finite recurrent de Sitter-patch interpretation. It does not prove that the real universe remains de Sitter for this duration, that the global universe is finite, or that a cosmological causal patch satisfies the microscopic assumptions of a finite quantum recurrence theorem.")

add_heading("6. Official DESI DR2 posterior propagation", 1)
add_para("The primary propagation no longer samples a fitted Gaussian. It applies the de Sitter branch calculation to the official weighted DESI DR2+CMB Cobaya posterior. A Gaussian approximation is retained only as a reproducibility cross-check.")
add_table(
    ["Quantity", "Median", "95% posterior interval"],
    [
        ["H0 (km s^-1 Mpc^-1)", f"{post.loc['H0','median']:.4f}", f"[{post.loc['H0','q2.5']:.4f}, {post.loc['H0','q97.5']:.4f}]"],
        ["Omega_m", f"{post.loc['omega_m','median']:.6f}", f"[{post.loc['omega_m','q2.5']:.6f}, {post.loc['omega_m','q97.5']:.6f}]"],
        ["H_Lambda^-1 (Gyr)", f"{post.loc['horizon_time_Gyr','median']:.4f}", f"[{post.loc['horizon_time_Gyr','q2.5']:.4f}, {post.loc['horizon_time_Gyr','q97.5']:.4f}]"],
        ["S_dS/k_B", sci(post.loc['S_dS_over_kB','median'],5), f"[{sci(post.loc['S_dS_over_kB','q2.5'],5)}, {sci(post.loc['S_dS_over_kB','q97.5'],5)}]"],
        ["log10(t_thermo/yr)", sci(post.loc['log10_t_thermo_rec_years','median'],5), f"[{sci(post.loc['log10_t_thermo_rec_years','q2.5'],5)}, {sci(post.loc['log10_t_thermo_rec_years','q97.5'],5)}]"],
        ["log10 log10(t_thermo/yr)", f"{post.loc['log10log10_t_thermo_rec_years','median']:.5f}", f"[{post.loc['log10log10_t_thermo_rec_years','q2.5']:.5f}, {post.loc['log10log10_t_thermo_rec_years','q97.5']:.5f}]"],
    ],
)
add_figure("de-sitter.png", "Figure 1. Weighted official DESI DR2+CMB posterior propagated to the compact double-log coordinate of the conditional de Sitter thermodynamic recurrence scale.")
add_para("The result is a scale-separation statement: current uncertainty in H0 and Omega_m shifts the enormous thermodynamic exponent only modestly, while a change in the far-future branch can remove the recurrence conclusion entirely.")

add_heading("7. Quantum epsilon recurrence is a different clock", 1)
add_para("Gupta and Short define finite-system quantum recurrence using trace distance. For a finite-dimensional unitary system with a finite discrete Hamiltonian spectrum, their continuous-time theorem gives an upper bound on the time at which all states return within epsilon, provided at least one state has moved farther than epsilon at an earlier time.")
add_equation("T(rho(t_r), rho(0)) <= epsilon")
add_equation("t_r <= [2 pi hbar / (E_max-E_min)] [2 ceil(pi/epsilon)]^(d-2)")
add_para("This theorem is rigorous for the stated finite-system assumptions, but CRPS does not substitute de Sitter entropy into it. A cosmological numerical application would require, at minimum, a justified finite Hamiltonian description, the number d of distinct energy eigenvalues, and the energy span E_max-E_min. Those microscopic inputs are not identified by current cosmological data.")
add_table(
    ["Metric", "epsilon", "Cosmological numerical value", "Reason"],
    [["Trace distance", e, "Not identified", "Finite discrete spectrum and energy span are unknown"] for e in [0.1, 0.01, 0.001]],
)
add_para("The clean conclusion is therefore two-clock, not one-clock: CRPS can compute a conditional thermodynamic de Sitter scale, while a cosmological epsilon-recurrence time remains uncalculated pending microscopic input.")

add_heading("8. Metastable de Sitter: recurrence versus decay", 1)
add_para("If the universe approaches a de Sitter-like phase but the vacuum is metastable, recurrence competes with vacuum decay. Under constant independent hazards:")
add_equation("P(recurrence before decay) = tau_decay / (tau_decay + t_rec)")
add_para("A vacuum lifetime that is enormous by astrophysical standards can still be negligible compared with an entropy-exponential recurrence scale. CRPS therefore performs competing-hazard calculations in logarithmic space. It does not assign a decay lifetime because present cosmological data do not measure the relevant microscopic vacuum-decay rate.")

add_heading("9. Finite-lifetime phantom scenarios", 1)
add_para("A constant equation of state w<-1 produces a finite future proper time in the standard late-time approximation:")
add_equation("Delta t ~= 2 / [3 |1+w| H0 sqrt(Omega_DE)]")
add_table(
    ["Constant w", "Approximate remaining lifetime", "Status"],
    [[f"{r.w:.2f}", f"{r.remaining_Gyr:,.1f} Gyr", "Scenario example only"] for r in bigrip.itertuples()],
)
add_figure("big-rip.png", "Figure 2. Constant-w phantom scenario lifetimes at the CPS reference point. These are not extrapolations of the DESI w0-wa fit to the infinite future.")
add_para("These examples illustrate why a finite future endpoint and an eternal de Sitter recurrence branch are qualitatively different hypotheses. Current evidence for evolving dark energy does not by itself determine a Big Rip or any other specific endpoint.")

add_heading("10. Recollapse, bounce, and cyclic futures", 1)
add_para("A Big Crunch does not by itself imply recurrence. Contraction is not a time-reversal operator, and a bounce is not automatically a reset. A genuinely recurrent cyclic model needs an explicit cycle-to-cycle map whose dynamics return to, or approach, a prior state under a declared metric.")
add_equation("x_{n+1} = F(x_n)")
add_para("Temporal recurrence then becomes a property of iterates of F, not a consequence of the word 'cycle'. Dissipation, entropy production, particle creation, phase changes, or hidden degrees of freedom can prevent state return even if the scale factor oscillates.")

add_heading("11. Probability result: what is and is not identified", 1)
add_para("If the branch weights and the recurrence-theory probabilities are allowed to vary over everything consistent with current observations, the marginalized recurrence probability can range from zero to one. The current-data-only result is therefore:")
add_equation("P(recurrence | current observations, unrestricted future theory) in [0,1]")
add_para("This is the formal consequence of an underidentified problem, not an informative probability estimate. Observations constrain the recent and intermediate expansion history; they do not yet identify the full dark-energy potential, vacuum lifetime, microscopic de Sitter state space, or a future cycle map. The full [0,1] interval means the target is non-identified under unrestricted future theory; it does not assign equal plausibility to every value.")

add_heading("12. Bayesian model-family weights", 1)
add_para("A probability-like quantity can be introduced at the model-family level if the model set and prior odds are declared. CRPS adopts, as a transparent baseline, the 2026 Bayesian reanalysis by Ong, Yallup, and Handley for DESI DR2 BAO + Planck CMB. For the restricted two-model comparison LambdaCDM versus w0-wa CDM, the published ln Bayes factor is -0.57 +/- 0.26 in favor of the dynamical model relative to LambdaCDM. With equal prior odds this maps to normalized posterior model weights.")
add_table(
    ["Data combination", "LambdaCDM weight", "Dynamical-DE weight", "Status"],
    [[r.dataset, f"{r.p_LambdaCDM*100:.1f}%", f"{r.p_dynamic_DE*100:.1f}%", r.status] for r in model_weights.itertuples()],
)
add_figure("model-weights.png", "Figure 3. Equal-prior posterior model weights implied by published Bayesian evidence values. The original DES-SN5YR comparison is shown only to expose calibration sensitivity and is not the CRPS baseline.")
add_para("These weights are conditional on the selected model set, prior odds, likelihood construction, and data combination. They are not direct probabilities that the universe ends in a Big Rip, Big Crunch, eternal de Sitter state, or cyclic future. DESI DR2 Results IV reports a 2.7 sigma frequentist preference for w0-wa over LambdaCDM for DESI+CMB, illustrating that frequentist fit significance and Bayesian model probability answer different questions.")

add_heading("13. Fate-probability decomposition", 1)
add_para("To make ultimate-fate uncertainty visible, CRPS decomposes the baseline top-level model weights into six named futures. The split within each model family is not observationally identified, so three explicit sensitivity maps are supplied rather than one claimed forecast. The maximum-entropy split is a neutral bookkeeping choice, not a physical principle.")
maxent = fate[fate.scenario == "maximum_entropy_split"]
label_map = {
    "stable_eternal_ds":"Stable eternal de Sitter",
    "metastable_ds":"Metastable de Sitter",
    "fading_dark_energy":"Fading dark energy",
    "phantom_finite_end":"Phantom finite end",
    "recollapse":"Recollapse",
    "cyclic_or_bounce":"Cyclic/bounce",
}
add_table(
    ["Future", "Illustrative weight", "Recurrence interpretation"],
    [[label_map[r.fate], f"{r.percent:.1f}%", {
        "stable_eternal_ds":"Conditional de Sitter recurrence route exists if finite recurrent-patch assumptions hold",
        "metastable_ds":"Competes with vacuum decay; decay lifetime unknown",
        "fading_dark_energy":"Standard de Sitter recurrence argument does not transfer automatically",
        "phantom_finite_end":"Finite endpoint; standard de Sitter thermodynamic route is effectively unavailable",
        "recollapse":"Crunch alone does not imply recurrence",
        "cyclic_or_bounce":"Potentially recurrent only if the cycle map returns states",
    }[r.fate]] for r in maxent.itertuples()],
)
add_figure("fate-sensitivity.png", "Figure 4. Fate-weight sensitivity using the Bayesian top-level model-family weights plus three declared within-family theory splits. These are model-dependent scenario weights, not measured ultimate-fate probabilities.")
add_para("Under the maximum-entropy bookkeeping split, the numerical weights are about 31.9% stable eternal de Sitter, 31.9% metastable de Sitter, and 9.0% each for fading dark energy, a phantom finite end, recollapse, and a cyclic/bounce future. The study does not privilege those six numbers as truth; their purpose is to show exactly how the answer moves when unresolved theory assumptions are changed.")

add_heading("14. Prior-conditioned recurrence sensitivity", 1)
add_para("For illustration only, CRPS retains three prior scenarios. Each scenario supplies both weights over future branches and a theoretical support parameter for the recurrent finite-patch assumptions conditional on eternal de Sitter. These are not observational posterior probabilities.")
add_table(
    ["Prior scenario", "p(recurrent patch | eternal dS)", "Conditional recurrence interval"],
    [[r.prior_set, f"{r.p_recurrent_patch_given_eternal_ds:.2f}", f"[{r.conditional_lower_bound:.3f}, {r.conditional_upper_bound:.3f}]"] for r in prior.itertuples()],
)
add_figure("prior-intervals.png", "Figure 5. Prior-conditioned theoretical sensitivity. The intervals depend on supplied theory weights and must not be reported as measured recurrence probabilities.")

add_heading("15. Finite-horizon probabilities and numerical precision", 1)
add_para("For a stationary rare-event model with mean wait t_thermo, P(T)=1-exp(-T/t_thermo). For T much smaller than t_thermo, log10 P is approximately log10 T - log10 t_thermo. A naive ordinary floating-point subtraction can cause many different horizons to print identically because their corrections are microscopic compared with a 10^122-scale exponent; the implementation therefore preserves the subtraction in a numerically stable representation.")
add_para("The study stores the rare-event subtraction with high-precision decimal arithmetic and also records the horizon exponent as a fraction of the recurrence exponent. This fixes the representation problem without pretending that the physical probabilities are meaningfully large.")

add_heading("16. Entropy-deficit ansatz", 1)
add_para("CRPS retains an illustrative macrostate ansatz, Delta S=f S_dS and t~H^-1 exp(Delta S), only to visualize sensitivity to a chosen entropy deficit. The linear relationship seen after taking double logarithms is a mathematical consequence of the assumed equation. It is not empirical evidence, and there is no derived mapping here from trace-distance epsilon to f.")
add_figure("entropy.png", "Figure 6. Visualization of the assumed entropy-deficit scaling. This figure is illustrative and is not an epsilon-recurrence calculation.")

add_heading("17. What current observations actually tell us", 1)
add_table(
    ["Finding", "Defensible interpretation"],
    [
        ["Flat LambdaCDM permits a precise conditional de Sitter entropy calculation", "Useful branch calculation; not proof of the far future"],
        ["DESI posterior uncertainty barely changes the qualitative thermodynamic recurrence scale", "Parameter uncertainty is secondary"],
        ["Published Bayesian DESI DR2+CMB comparison gives roughly 64% LambdaCDM vs 36% w0-wa CDM under equal prior odds", "Useful model-family weight; not a fate probability"],
        ["Recent DESI analyses continue to test evolving dark energy", "Far-future extrapolation remains structurally uncertain"],
        ["Finite quantum recurrence theorems require explicit microscopic inputs", "Entropy alone is not a substitute for epsilon and spectrum"],
        ["Vacuum decay or finite future endpoints can preempt recurrence", "Ultimate-fate physics dominates"],
        ["A crunch or bounce does not automatically recreate the same state", "Cycle-map dynamics must be specified"],
        ["The theory-agnostic marginalized probability is [0,1]", "No unique current recurrence percentage is identified"],
    ],
)

add_heading("18. Limitations", 1)
add_bullet("The de Sitter entropy argument is branch-conditional and does not establish a finite Hilbert-space interpretation of quantum gravity in de Sitter space.")
add_bullet("The thermodynamic recurrence scale is a heuristic/thermodynamic scale, not a rigorous system-specific first-return distribution for cosmology.")
add_bullet("The finite-system epsilon-recurrence theorem is not numerically instantiated for the universe because the necessary Hamiltonian spectrum is unknown.")
add_bullet("The DESI posterior propagation assumes the flat-LambdaCDM branch for that calculation; extended dark-energy models are not extrapolated into the remote future without a physical potential.")
add_bullet("The Big Rip calculations are constant-w scenarios, not forecasts from the current w0-wa phenomenology.")
add_bullet("The Bayesian model-family weights are conditional on a restricted two-model set and equal prior odds; they do not exhaust dark-energy theory space.")
add_bullet("The fate-decomposition percentages combine data-informed top-level model weights with explicit within-family theory splits and therefore are not observationally measured ultimate-fate probabilities.")
add_bullet("The prior-conditioned probability intervals are sensitivity demonstrations. Their branch weights and recurrent-patch support parameters are theoretical inputs.")
add_bullet("Spatial duplicates elsewhere in an infinite or inflating spacetime are not counted as temporal recurrence of our causal history.")

add_heading("19. Update pathways", 1)
add_para("CRPS is designed to narrow only when new information actually identifies one of its currently free components. High-value updates include:")
add_bullet("new DESI full-survey dark-energy posterior products and independent expansion-history measurements;")
add_bullet("a fundamental dark-energy model with a controlled far-future potential rather than a phenomenological fit alone;")
add_bullet("a quantitative vacuum-decay lifetime applicable to our vacuum;")
add_bullet("a quantum-gravity result establishing or rejecting an effectively finite de Sitter state space;")
add_bullet("a microscopic cosmological Hamiltonian or effective spectrum sufficient to evaluate an epsilon-recurrence theorem;")
add_bullet("a concrete cyclic/bounce model with an explicit return map and recurrence metric.")

add_heading("20. Conclusion", 1)
add_para("The universe may admit recurrence under some physically viable futures, but current observations do not tell us that it will recur. Under an eternal, stable, positive-Lambda future and the conventional finite recurrent causal-patch interpretation, the thermodynamic return scale is roughly 10^(1.38 x 10^122) years in the parameter region favored by current flat-LambdaCDM fits. That number is conditional, not a countdown.")
add_para(f"The strongest result of CRPS is the uncertainty decomposition. At the data-informed model-family level, one published equal-prior Bayesian comparison assigns about {model_weights.iloc[0].p_LambdaCDM*100:.1f}% weight to LambdaCDM and {model_weights.iloc[0].p_dynamic_DE*100:.1f}% to w0-wa CDM for DESI DR2+CMB. But those weights still do not determine vacuum stability or the remote fate of dynamical dark energy. Observational uncertainty in H0 and Omega_m is already small compared with those structural uncertainties. A defensible study therefore reports model-family weights, conditional scales, and explicit theory sensitivity instead of a fabricated universal recurrence percentage.")
add_para("Accordingly, the current theory-agnostic marginalized recurrence probability remains the full [0,1] logical bound, a non-identifiability result. Narrower values become meaningful only after the assumptions that produce them are supplied and defended.")

add_heading("References", 1)
refs = [
    "1. DESI Collaboration / M. Abdul Karim et al. (2025). DESI DR2 Results II: Measurements of Baryon Acoustic Oscillations and Cosmological Constraints. Physical Review D 112, 083515. arXiv:2503.14738.",
    "2. DESI Collaboration (2025-2026). DESI DR2 BAO cosmology products and public Cobaya chains. DESI Data portal.",
    "3. DESI Collaboration (2026). DESI DR2 Results IV: Alcock-Paczynski Measurements from the Lyman Alpha Forest and Cosmological Constraints. arXiv:2607.27410v3.",
    "4. Ong, D. D. Y., Yallup, D., & Handley, W. (2026). The Bayesian view of DESI DR2 with unimpeded: Evidence and tension in a combined analysis with CMB and supernovae across cosmological models. arXiv:2603.05472.",
    "5. Gupta, C., & Short, A. J. (2026). Recurrence Time for Finite Quantum Systems. arXiv:2604.14995v2. Journal metadata omitted here because the arXiv-listed journal DOI currently conflicts with the APS record.",
    "6. Gibbons, G. W., & Hawking, S. W. (1977). Cosmological event horizons, thermodynamics, and particle creation. Physical Review D 15, 2738.",
    "7. Dyson, L., Kleban, M., & Susskind, L. (2002). Disturbing Implications of a Cosmological Constant. JHEP 10 (2002) 011. arXiv:hep-th/0208013.",
    "8. Goheer, N., Kleban, M., & Susskind, L. (2003). The Trouble with de Sitter Space. JHEP 07 (2003) 056. arXiv:hep-th/0212209.",
    "9. Caldwell, R. R., Kamionkowski, M., & Weinberg, N. N. (2003). Phantom Energy and Cosmic Doomsday. Physical Review Letters 91, 071301. arXiv:astro-ph/0302506.",
    "10. Sundy, S. A. (2026). Cosmic Coordinate. Declared reference artifact.",
    "11. Sundy, S. A. (2026). Recurrence Dynamics Study . Terminology reference only; not incorporated into CRPS.",
]
for r in refs:
    add_para(r)

add_heading("Appendix A. Reproducibility and provenance", 1)
add_para("The repository includes the exact configuration, source code, compressed DESI posterior projection, source-chain hashes, numerical outputs, figures, and 38 automated tests. `python run.py` regenerates all numerical CSV outputs and figures from `config.json`. `python -m pytest` validates the core equations, weighted posterior machinery, input validation, probability bounds, finite-horizon precision behavior, and bundled observational projection.")
add_para("The DESI projection contains 59,891 source rows with summed Cobaya weight 169,444. The four source chain SHA-256 hashes and official directory are recorded in `data/desi-source.md`; `scripts/fetch_desi.py` independently rebuilds the projection after verifying those hashes.")

add_heading("Appendix B. Interpretation rules", 1)
rules = [
    "Never report a near recurrence as exact recurrence.",
    "Never call the de Sitter entropy-exponential thermodynamic scale an epsilon-recurrence time without a declared microscopic model, metric, and tolerance.",
    "Never interpret a spatial duplicate as temporal recurrence.",
    "Never treat a w0-wa fit as a guaranteed far-future law without an explicit physical extrapolation model.",
    "Never replace an unidentified probability with 0.5 merely because it is unknown.",
    "Never present the prior-sensitivity scenarios as observational posterior probabilities.",
    "Always separate observational parameter uncertainty from far-future model and recurrence-theory uncertainty.",
]
for x in rules:
    add_bullet(x)

add_heading("Appendix C. Data acknowledgment", 1)
add_para("This study uses public DESI cosmological posterior data. CRPS is independent work and is not an official DESI analysis or endorsement. The repository directs downstream users to the current DESI data-license and acknowledgment requirements in `ACKNOWLEDGMENTS.md` and `data/desi-source.md`.")

# Footer page number field.
for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("the Cosmological Recurrence Study  |  ")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

md_path = PAPER / "paper.md"
md_path.write_text("\n".join(md), encoding="utf-8")
docx_path = PAPER / "cosmological_recurrence_probability.docx"
doc.save(docx_path)
print(docx_path)
print(md_path)
